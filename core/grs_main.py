import math
import random
import sqlite3
import turtle
from tkinter import *
from tkinter import scrolledtext
from tkinter.ttk import Combobox
import tkinter.ttk as ttk
import pandas as pd
import matplotlib.pyplot as plt

import CoolProp.CoolProp as CP
import docx

import grs_stat
from core.entity.Gas import Gas
from core.entity.Pipeline import Pipeline
from core.entity.Vessel import Vessel

component_list = ['Methane', 'Ethane', 'Propane', 'Isobutane', 'Butane',
                 'Isopentane', 'Pentane', 'Hexane', 'Oxygen',
                 'Nitrogen', 'CarbonDioxide']
component_list_for_query = ['methane', 'ethane', 'propane', 'isobutane', 'butane',
                 'isopentane', 'pentane', 'hexane', 'oxygen',
                 'nitrogen', 'carbon_dioxide']

"""функция выполнения sql запросов, расположенных в текстовом файле"""
def sql_query(file_name, *args):
    connection = sqlite3.connect('grs_database.db')
    cur = connection.cursor()
    with open(file_name, 'r', encoding='utf-8') as f:
        queries = f.read().split(';')
    for query in queries:
        cur.execute(query, args)
    connection.commit()


def get_grs_name():
    connection = sqlite3.connect('grs_database.db')
    cur = connection.cursor()
    cur.execute("""SELECT name_grs FROM grs""")
    grs_list = [grs[0] for grs in cur.fetchall()]
    connection.close()
    return grs_list


"""выбирает данные указанного столбца указанной таблицы для указанной грс, написано плохо,
так как привязано к grs_id, а он не везде в таблицах есть"""
def get_column_data(column_name, table, name_grs):
    connection = sqlite3.connect('grs_database.db')
    cur = connection.cursor()
    cur.execute(f"""SELECT {column_name} FROM {table} WHERE grs_id = 
    (SELECT grs_id FROM grs WHERE name_grs = '{name_grs}')""")
    data_list = [data[0] for data in cur.fetchall() if data[0] is not None]
    if data_list == []:
        data_list = ['Данные отсутствуют']
    connection.close()
    return data_list


"""эта функция получает состав газа ГРС с заданным именем, параметр нормализации отвечает за приведение
состава к такому, чтобы сумма компонентов была крайне близка к 1"""
def get_composition(name_grs, normolize=True):
    connection = sqlite3.connect('grs_database.db')
    cur = connection.cursor()
    component_list_query = ', '.join(component for component in component_list_for_query)
    cur.execute(f"""SELECT {component_list_query} FROM composition
                    WHERE grs_id = (SELECT grs_id FROM grs 
                    WHERE name_grs = '{name_grs}');""")
    values = cur.fetchone()
    if normolize == True:
        while abs(sum(values) - 1) > 10e-10:
            values = [component + ((1 - sum(values)) * component) for component in values]
    composition = dict(zip(component_list, values))
    connection.close()
    return composition

"""список файлов, содержащих запросы для создания и наполнения таблиц данными"""
file_list = ['create_grs_table',
             'insert_grs_data',
             'insert_grs_compositions',
             'equipment_list']


"""выполнение всех запросов, находящихся во всех файлах"""
def reboot_data():
    for file in file_list:
        file +='.txt'
        sql_query(file)
        print(f'обработка таблицы {file}, прошла успешно!')
    print(f'создание таблиц и их наполнение прошло успешно!')


"""Так как мне лень, эта функция переводит нормализованный состав газа в удобоваримый для хайсиса вариант"""
def composition_to_hysys(name_grs):
    composition = [value for value in get_composition(name_grs).values()]
    file_name = f'{name_grs} состав для Hysys.txt'
    with open(file_name, 'w', encoding='utf-8') as f:
        for value in composition:
            f.write(f'{value}\n')


# ОКНО --- Работа с базой данных ГРС
def grs_data():
    grs_data_root = Tk()
    grs_data_root.title('Работа с базой данных ГРС')
    grs_data_root.geometry('1920x1080')

    txtwin = scrolledtext.ScrolledText(grs_data_root, width=150, height=15, bg="darkgreen", fg='white')
    txtwin.pack()

    text = Text(grs_data_root, width=100, height=30)
    text.pack(pady=10)

    lbl = Label(grs_data_root, text='файл с запросами', font=("Arial Bold", 10))
    lbl.pack(pady=10)

    entry = Entry(grs_data_root, width=20)
    entry.pack(pady=5)


    def make_query():
        connection = sqlite3.connect('grs_database.db')
        cur = connection.cursor()
        if entry.get() != '':
            path = entry.get() + '.txt'
            sql_query(path)
            print(f'запрос из файла {path} успешно выполнен')
        else:
            query = text.get(1.0, END)
            cur.execute(query)
            result = cur.fetchall()
            txtwin.insert(INSERT, result)
            connection.commit()
            print('Запрос выполнен успешно!')
            txtwin.insert(INSERT, '\nЗапрос выполнен успешно!\n')
    def clear():
        txtwin.delete(1.0, END)


    query_btn = Button(grs_data_root, text='Сделать запрос', command = make_query, bg='blue', fg='white')
    query_btn.pack()

    clear_btn = Button(grs_data_root, text='Почистить', command=clear, bg='green', fg='white')
    clear_btn.pack()

    btn = Button(grs_data_root, text='Пересоздать данные', command=reboot_data, bg='red', fg='white')
    btn.pack()


    grs_data_root.mainloop()

"""рисует схему трубопроводов по заданным в базе данных координатам"""
def pipe_draw():

    """рисует линию если надо, если не надо, то не рисует, всё логично"""
    def draw_arrow(x, y, draw_line=True, color='Blue'):
        t.color(color)
        if draw_line:
            t.goto(x, y)
        else:
            t.penup()
            t.goto(x, y)
            t.pendown()


    pipe_id_list = []
    x_coord_list = []
    y_coord_list = []

    connection = sqlite3.connect('grs_database.db')
    cur = connection.cursor()
    cur.execute("""SELECT * FROM knots""")
    for i, j, k, n in cur.fetchall():
        pipe_id_list.append(j)
        x_coord_list.append(k)
        y_coord_list.append(n)

    color_list = ['Blue', 'Red', 'Green', 'Purple']

    t = turtle.Turtle()
    t.shape('arrow')
    color = 'Blue'
    t.color(color)
    t.speed(10)
    t.penup()
    start_pos = (-600, 400)
    t.setpos(*start_pos)
    t.pendown()
    t.pensize(5)
    print(max(y_coord_list, key=abs))
    y_turtle_scale = abs(800 / max(y_coord_list, key=abs))
    x_turtle_scale = abs(1200 / max(x_coord_list, key=abs))


    for x, y, color in zip(x_coord_list, y_coord_list, pipe_id_list):
        x = x * x_turtle_scale + start_pos[0]
        y = y * y_turtle_scale + start_pos[1]
        colour = color_list[color-1]
        draw_arrow(x, y, True, colour)


# ОКНО --- Расчёт трубопроводов
def pipe_wall_root():
    pipe_root = Tk()
    pipe_root.geometry("1920x1080")
    pipe_root.title('Расчёт трубопроводов')

    connection = sqlite3.connect('grs_database.db')

    cur = connection.cursor()

    cur.execute("""SELECT class FROM class_mat;""")
    class_list = [cathegory[0] for cathegory in cur.fetchall()]

    txtwin = scrolledtext.ScrolledText(pipe_root, width=170, height=15, bg="darkgreen", fg='white')
    txtwin.pack()

    text = Text(pipe_root, width=100, height=15)
    text.pack(pady=10)


    entry_list = []
    lbl_list = []

    blocks_data = ['Узел переключения', 'Узел очистки', 'Узел предотвращения гидратообразования',
                   'Узел смешения', 'Узел учёта', 'Узел обводной линии', 'Узел одоризации']

    combo_blocks = Combobox(pipe_root, width=50)
    combo_blocks['values'] = blocks_data
    block_lbl = Label(pipe_root, text='Узел')
    block_lbl.pack()
    combo_blocks.pack()
    pipe_data = ['DN', 'PN', 'внешний диаметр', 'толщина стенки', 'материал', 'категория прочности',
                     'длина']

    lbl = Label(pipe_root, text='файл с запросами', font=("Arial Bold", 10))
    lbl.pack()

    entry = Entry(pipe_root, width=20)
    entry.pack()
    grs_lbl = Label(pipe_root, text='Наименование ГРС')
    grs_lbl.pack()
    combo = Combobox(pipe_root)
    combo['values'] = get_grs_name()
    combo.current(0)
    combo.pack()

    pipe_root.mainloop()


def pipe_shell_thickness(name_grs):
    connection = sqlite3.connect('grs_database.db')
    cur = connection.cursor()

    """запрос в базу данных по характеристикам материалов"""
    cur.execute("""SELECT temp_resistance_min
     FROM pipes as p INNER JOIN class_mat as cm 
     ON p.strength_category = cm.class 
     LIMIT 1;
    """)
    R1n = cur.fetchone()[0]
    print(f'минимальное временное сопротивление материала труб составило {R1n}')
    k1 = 1.55
    kn = 1.1
    n = 1.1
    m = 0.66
    R1 = round(R1n * m / k1 / kn, 2)

    print(f'расчётная величина R1 - - - > {R1}')
    cur.execute(f"""SELECT inlet_design_pressure FROM grs WHERE name_grs = '{name_grs}';""")
    design_pressure = cur.fetchone()[0]
    print(f'Для ГРС {name_grs} проектное давление на входе задано как {design_pressure}')

    """этот запрос вытягивает из базы данных выбранной ГРС все трубы, что там есть и """
    cur.execute(f"""SELECT DISTINCT external_diameter 
    FROM pipes 
    WHERE grs_id = (
        SELECT grs_id 
        FROM grs 
        WHERE name_grs = '{name_grs}') 
    ORDER BY 1 DESC;""")
    diameters_list = [diameter[0] for diameter in cur.fetchall()]

    calc_thickness_list = []

    for diameter in diameters_list:
        calc_thickness = n * design_pressure * diameter / 2 / (R1 + n * design_pressure)
        calc_thickness_list.append(round(calc_thickness, 2))
        print(f'Для трубы {diameter} мм расчётная толщина стенки ----> {round(calc_thickness, 2)}')


# ОКНО --- Расчёт скорости в трубопроводах
def veloсity_calc():
    velocity_root = Tk()
    velocity_root.geometry("1920x1080")
    velocity_root.title('Расчёт скорости в трубопроводах')
    combo = Combobox(velocity_root)
    combo['values'] = get_grs_name()
    combo.current(0)
    combo.grid(column=100, row=0)

    manual_input = ['Давление, МПа', 'Температура, °С', 'Расход', 'Диаметр Ду, мм']

    lbl = []
    txt = []

    for i in range(len(manual_input)):
        lbl.append(Label(velocity_root, text=manual_input[i], font=("Arial Bold", 14)))
        lbl[i].grid(column=500, row=100 * i)

        entry = Entry(velocity_root, width=10, justify='center')
        entry.grid(column=600, row=100 * i)
        txt.append(entry)
    txtwin = scrolledtext.ScrolledText(velocity_root, width=80, height=40, bg="darkgreen", fg='white')
    txtwin.grid(column=10, row=500)

    interface_buttons = ['Скорость газа',
                         'Плотность газа',
                         'Подбор диаметра',
                         'Очистить окно вывода',
                         'Рассчитать всю ГРС']

    def calc():
        """Рассчитывает скорость газа в трубопроводе"""

        chois = get_composition(combo.get())

        result = (f'Запуск расчёта ГРС {format(combo.get())}\n')
        txtwin.insert(INSERT, result)
        res = [format(data.get()) for data in txt]
        if '' in res:
            result = 'Данные введены не полностью!\n'
            txtwin.insert(INSERT, result)
            return
        pressure = float(res[0].replace(',', '.'))
        temperature = float(res[1].replace(',', '.'))
        rate = float(res[2].replace(',', '.'))
        internal_diameter = float(res[3].replace(',', '.'))

        pipe = Pipeline(internal_diameter, 0, 0)
        gas = Gas(chois, temperature, pressure, rate)
        velosity = gas.actual_rate / pipe.area / 3600
        capacity = round(25 * rate / velosity)
        result = (f'При давлении {pressure} МПа\n'
                  f'температуре {temperature} °C\n'
                  f'расходе {rate}  ст м3/ч\n'
                  f'в трубопроводе Ду {internal_diameter} мм\n'
                  f'Скорость составила {round(velosity, 2)} м/с \n'
                  f'Пропускная способность основных линий {capacity} ст. м3/ч\n'
                  f'Для обводных линий {capacity * 50 / 25} ст. м3/ч\n')
        txtwin.insert(INSERT, result)
        return velosity


    def density():
        """выводит на экран плотность газа при рабочих,нормальных, стандартных условиях, актуальный расход газа"""
        chois = get_composition(combo.get())

        res = [format(data.get()) for data in txt]
        pressure = float(res[0].replace(',', '.'))
        temperature = float(res[1].replace(',', '.'))
        rate = float(res[2].replace(',', '.'))
        gas = Gas(chois, temperature, pressure, rate)
        result = (
            f'Газ ГРС {format(combo.get())}.\nПлотность при рабочих условиях {round(gas.actual_density, 2)} кг/м3 \nПлотность при нормальных условиях {round(gas.normal_density, 4)} кг/м3 \n'
            f'Плотность при стандартных условиях {round(gas.standard_density, 4)} кг/м3 \n'
            f'Расход газа при рабочих условиях {round(gas.actual_rate, 2)} м3/ч\n'
            f'Молярная масса газа {round(gas.molecular_mass, 2)} кг/кмоль\n')
        txtwin.insert(INSERT, result)

    def pipecheck():
        """определяет требуемый диаметр трубопровода чтобы скорость была в порядке"""
        chois = get_composition(combo.get())
        velosity_limit = 25
        result = (f'Запуск расчёта ГРС {format(combo.get())}\n')
        txtwin.insert(INSERT, result)
        res = [format(data.get()) for data in txt]
        pressure = float(res[0].replace(',', '.'))
        temperature = float(res[1].replace(',', '.'))
        rate = float(res[2].replace(',', '.'))
        gas = Gas(chois, temperature, pressure, rate)
        actual_flow = gas.standard_density / gas.actual_density * rate
        internal_diameter = (actual_flow * 4 / 3600 / velosity_limit / math.pi) ** 0.5 * 1000
        result = (f'Минимальный диаметр прохода {round(internal_diameter, 2)}, мм \n')
        txtwin.insert(INSERT, result)

    def clean():
        txtwin.delete(1.0, END)  # чистит окно вывода

    def full_calc():
        name_grs = combo.get()
        connection = sqlite3.connect('grs_database.db')
        cur = connection.cursor()
        cur.execute(f"""SELECT pipe_diameter, wall FROM tvps
        WHERE block_id = (SELECT block_id FROM blocks WHERE grs_id = (SELECT grs_id FROM grs 
        WHERE name_grs  = '{format(combo.get())}')
        """)

        print(f'начинаю расчёт ТВПС ГРС {name_grs}')
        gas_velocity = []
        diameter_list = []
        for ex_diameter, wall in cur.fetchall():
            diameter_list.append(float(ex_diameter) - 2 * float(wall))

        if len(diameter_list) > 0:
            db_load = 'База данных загружена!\n'
        else:
            db_load = 'База не загружена!!!\n'
        txtwin.insert(INSERT, db_load)
        for i in range(len(diameter_list)):
            progress = round(i / len(diameter_list) * 100, 2)
            print(f'Рассчитано {progress} %')
            txt[3].insert(0, diameter_list[i])
            gas_velocity.append(calc())
            tvps = round(float(format(txt[2].get())) / calc() * 25)
            if gas_velocity[i] > 25:
                txtwin.insert(INSERT, 'ПРЕВЫШЕНИЕ СКОРОСТИ ГАЗА\n')
            txtwin.insert(INSERT, '\n')
            txt[3].delete(0, 'end')
        connection.close()

    calc_list = [calc, density, pipecheck, clean, full_calc]
    for i in range(len(interface_buttons)):
        btn = Button(velocity_root, text=interface_buttons[i], font=("Arial Bold", 10), bg="green", fg="black",
                     command=calc_list[i])
        btn.grid(column=800, row=400 + i * 100)

    velocity_root.mainloop()
    return

# ОКНО --- Расчёт ёмкости одоранта
def odorant_calc():
    odorant_root = Tk()
    odorant_root.geometry("1920x1080")
    odorant_root.title('Расчёт ёмкости одоранта')
    manual_input = ['Вместимость ёмкости', 'Расход газа в ст. м3']
    lbl = []
    txt = []
    for i in range(len(manual_input)):
        lbl.append(Label(odorant_root, text=manual_input[i], font=("Arial Bold", 14)))
        lbl[i].grid(column=500, row=100 * i)
        entry = Entry(odorant_root, width=10, justify='center')
        entry.grid(column=600, row=100 * i)
        txt.append(entry)

    txtwin = scrolledtext.ScrolledText(odorant_root, width=100, height=50, bg="darkgreen", fg='white')
    txtwin.grid(column=10, row=500)

    def calc():
        """рассчитывает запас одоранта и выводит на экран"""
        clean()
        res = [format(data.get()) for data in txt]
        if '' in res:
            result = 'Данные введены не полностью!\n'
            txtwin.insert(INSERT, result)
            return

        volume = float(res[0].replace(',', '.'))
        rate = float(res[1].replace(',', '.'))
        odorant_vessel = Vessel(0, 0, volume)

        if odorant_vessel.odorant_time(volume, rate) < 60:
            verdict = 'не'
        else:
            verdict = ''

        result = (f'Объём одоранта в ёмкости хранения после заправки составляет {volume} м3×0,85 = {volume * 0.85} м3.\n'
              f'При плотности одоранта 830 кг/м3 масса одоранта составит {volume * 0.85} м3×830 кг/м3 '
                  f'= {odorant_vessel.odorant_mass(volume)} кг.\n'
              f'При рассчитанном предельном расходе газа через выходной коллектор {rate} м3/ч\n'
              f'и нормативном расходе одоранта 16 г на 1000 ст. м3 (п. 9.7.2 СТО Газпром 2-3.5-051-2006)\n'
              f'такого количества хватит на {odorant_vessel.odorant_mass(volume)} кг/ '
                  f'{odorant_vessel.odorant_rate(rate)} кг/ч = '
              f'{odorant_vessel.odorant_mass(volume) / odorant_vessel.odorant_rate(rate)} ч '
                  f'≈ {odorant_vessel.odorant_time(volume, rate)} суток.\n'
              f'Следовательно, имеющаяся ёмкость хранения одоранта {verdict} обеспечивает пропускную способность\n'
              f'в соответствии с требованиями СТО Газпром 2-3.5-051-2006, пункт 9.7.6.\n')

        #result = f'Запаса одоранта хватит на {odorant_vessel.odorant_time(volume, rate)} дней\n'
        txtwin.insert(INSERT, result)


    def vessel_check():
        """рассчитывает необходимый запас одоранта"""
        rate = float(format(txt[1].get()).replace(',', '.'))
        request_volume = rate / 1000 * 0.016 * 60 * 24 / 830
        result = f'Минимальный требуемый объём ёмкости одоранта {round(request_volume, 2)} м3\n'
        txtwin.insert(INSERT, result)


    def clean():
        """чистит окно вывода"""
        txtwin.delete(1.0, END)

    interface_buttons = ['Запас', 'Требуемый объём', 'Очистить окно вывода']
    calc_list = [calc, vessel_check, clean]
    for i in range(len(interface_buttons)):
        btn = Button(odorant_root, text=interface_buttons[i], font=("Arial Bold", 10), bg="green", fg="black",
                     command=calc_list[i])
        btn.grid(column=800, row=400 + i * 100)  # добавляет кнопки интерфейса

    return


def valve_calc():
    def create_word():
        return
        mydoc = docx.Document()
        mydoc.add_paragraph('Добрый день! Исходные данные для расчёта пропускной способности клапана:')
        for i in range(len(manual_input)):
            mydoc.add_paragraph(manual_input[i] + 5 * '....' + format(txt[i].get()))
        mydoc.add_paragraph('Состав газа:')

        calc()  # выполняется расчёт
        for i, j in chois.items():
            mydoc.add_paragraph(i + '----' + str(round(j * 100, 3)) + '%')  # запись исходных данных для расчёта
        mydoc.add_paragraph('Результаты расчёта: ')
        mydoc.add_paragraph('Плотность газа при н.у. ' + str(round(gas.normal_density, 4)) + 'кг/м3')
        mydoc.add_paragraph(result)

        file_name = 'Результаты расчёта ' + format(combo.get()) + ' ' + format(txt[0].get()) + ' МПа---' + format(
            txt[1].get()) + ' МПа.docx'

        mydoc.save(file_name)



    def shift():
        if btn1.cget('text') == 'Расчёт Kv':
            lbl[3].config(text='Расход газа, ст. м3/ч')
            btn1.config(text='Расчёт пропускной способности')
        else:
            lbl[3].config(text='Kv')
            btn1.config(text='Расчёт Kv')

    def clean():
        txtwin.delete(1.0, END)  # чистит окно вывода

    def calc(event=None):
        global result
        clean()

        design_pressure = 5.39

        connection = sqlite3.connect('grs_database.db')
        cur = connection.cursor()

        res = [format(data.get()) for data in txt]
        if '' in res:
            result = ('!!!Введите исходные данные!!!\n')
            txtwin.insert(INSERT, result)
            return

        p1 = float(res[0].replace(',', '.'))
        p2 = float(res[1].replace(',', '.'))
        if p2 >= p1:
            result = 'Обратный перепад давления!\n' \
                     'Проверьте исходные данные!'
            txtwin.insert(INSERT, result)
            return
        temperature = float(res[2].replace(',', '.'))

        def valve_calculation(normal_density, temperature, p1, p2):

            # этот кусок кода нужен, чтобы найти плотность газа при нормальных условиях
            p1 = (p1 + 0.101325) * 1000000 / 98100
            p2 = (p2 + 0.101325) * 1000000 / 98100
            dp = p1 - p2
            temperature += 273.15
            if btn1.cget('text') == 'Расчёт пропускной способности':
                g = normal_density * float(res[3].replace(',', '.'))
                if dp < p1 / 2:
                    kv = g / 529 * (temperature / dp / p2 / normal_density) ** 0.5
                else:
                    kv = g / 265 / p1 * (temperature / normal_density) ** 0.5
                cur.execute("""SELECT name, link FROM pressure_regulators
                                WHERE Kv>(?);
                """, (kv*1.2,))
                pressure_regulator = cur.fetchall()
                print(f'Регуляторы {pressure_regulator}')
                result = f'Требуемый Kv  {round(kv, 1)} м3/ч\n' \
                         f'с учётом запаса {round(kv * 1.2, 1)} - {round(kv * 1.5, 1)}\n' \
                         f'\nРегуляторы\n{pressure_regulator}\n'

            elif btn1.cget('text') == 'Расчёт Kv':
                kv = float(res[3].replace(',', '.'))
                args = [temperature, kv, dp, p1, p2, normal_density]
                if long_mode is True:
                    result = f'{round(standard_conditions(*args)/1.2, 1)}\n'
                else:
                    result = f'Пропускная способность {round(mass_rate(*args), 1)} кг/ч \n' \
                         f'{round(normal_conditions(*args), 1)} нм3/ч \n' \
                         f'{round(standard_conditions(*args), 1)} ст. м3/ч\n' \
                         f'С учётом запаса - {round(standard_conditions(*args)/1.5, 1)} - ' \
                         f'{round(standard_conditions(*args)/1.2, 1)} ст. м3/ч\n'


            txtwin.insert(INSERT, result)


        long_mode = False
        if long_mode is True:
            while p1 <= design_pressure:
                valve_calculation(normal_density, temperature, p1, p2)
        else:
            valve_calculation(normal_density, temperature, p1, p2)


    def mass_rate(temperature, kv, dp, p1, p2, gas_density):
        if dp < p1 / 2:
            return (529 / temperature) * kv * (dp * p2 * gas_density * temperature) ** 0.5
        else:
            return 265 * p1 * kv * (gas_density / temperature) ** 0.5


    def normal_conditions(temperature, kv, dp, p1, p2, gas_density):
        return mass_rate(temperature, kv, dp, p1, p2, gas_density) / gas_density


    def standard_conditions(temperature, kv, dp, p1, p2, gas_density):
        return normal_conditions(temperature, kv, dp, p1, p2, gas_density) * 293.15 / 273.15


    valve_root = Tk()
    valve_root.geometry("1920x1080")
    valve_root.title('Расчёт пропускной способности клапанов ')

    manual_input = ['Давление до клапана, МПа', 'Давление после клапана, МПа', 'Температура, °С', 'Kv, м3/ч']

    lbl = []
    txt = []

    for i in range(len(manual_input)):
        lbl.append(Label(valve_root, text=manual_input[i], font=("Arial Bold", 14)))
        lbl[i].grid(column=100, row=700 + 100 * i)

        entry = Entry(valve_root, width=10, justify='center')
        entry.grid(column=110, row=700 + 100 * i)
        txt.append(entry)

    txtwin = scrolledtext.ScrolledText(valve_root, width=80, height=30, bg="darkgreen", fg='white')
    txtwin.grid(column=10, row=10)


    combo = Combobox(valve_root)
    combo['values'] = get_grs_name()
    combo.current(0)
    combo.grid(column=300, row=0)
    """при каждой смене ГРС, будет происходить инициализация нового газа, чтобы не проводить её тысячи раз"""
    def gas_init(event):
        gas = Gas(get_composition(combo.get()), 0, 0)
        return gas.normal_density
    normal_density = Gas(get_composition(combo.get()), 0, 0).normal_density
    combo.bind("<<ComboboxSelected>>", gas_init)

    btn = Button(valve_root, text='Выполнить расчёт', font=("Arial Bold", 10), bg="green", fg="black", command=calc)
    btn.grid(column=600, row=100)
    btn1 = Button(valve_root, text='Расчёт Kv', font=("Arial Bold", 10), bg="green", fg="black", command=shift)
    btn1.grid(column=600, row=200)
    #btn2 = Button(valve_root, text='Полный расчёт', font=("Arial Bold", 10), bg="blue", fg="white", command=long_calc)
    #btn2.grid(column=700, row=200)

    valve_root.bind('<Return>', calc)

    valve_root.mainloop()
    return


# ОКНО --- Расчёт предохранительных клапанов ГРС
"""код рассчитывает пропускную спобосность СППК для ГРС ориентировочно"""
def ppk_calc():

    def shift():
        if btn1.cget('text') == 'Расчёт седла':
            lbl[3].config(text='Площадь седла, мм2')
            btn1.config(text='Площадь седла')
        else:
            lbl[3].config(text='Расход газа, ст. м3/ч')
            btn1.config(text='Расчёт седла')


    def clean():
        txtwin.delete(1.0, END)  # чистит окно вывода

    def calc(event=None):
        global result
        clean()


        res = [format(data.get()) for data in txt]
        if '' in res:
            result = ('!!!Введите исходные данные!!!\n')
            txtwin.insert(INSERT, result)
            return

        chois = get_composition(combo.get())

        p1 = float(res[0].replace(',', '.'))            #давление до клапана
        temperature = float(res[1].replace(',', '.'))   #температура до клапана
        alpha = float(res[2].replace(',', '.'))         # коэфф. расхода
        rate = float(res[3].replace(',', '.'))          #расход газа в ст. м3/ч

        gas = Gas(chois, temperature, p1, rate)

        temperature += 273.15

        # этот кусок кода нужен, чтобы найти плотность газа при нормальных условиях

        print(f'Плотность газа составляет {round(gas.normal_density, 4)} кг/м3')
        k = 1.30 #адиабат. коэффициент для метана
        b3 = 1.59 * ((k / (k + 1)) ** 0.5) * ((2 / (k + 1)) ** (1 / (k - 1))) #расчёт b3
        if btn1.cget('text') == 'Расчёт седла':
            f_area = gas.standard_density * float(res[3].replace(',', '.')) / alpha / b3 / 3.16 / (((p1 + 0.1) * gas.actual_density)**0.5)
            result = f'Требуемая площадь седла  {round(f_area, 1)} мм2\n'
        elif btn1.cget('text') == 'Площадь седла':
            f_area = float(res[3].replace(',', '.'))
            args = [gas, alpha, f_area, p1]
            result = f'Пропускная способность ППК с седлом {f_area} мм2\n'\
                     f' {round(mass_rate(*args), 1)} кг/ч \n' \
                     f'{round(normal_conditions(*args), 1)} нм3/ч \n' \
                     f'{round(standard_conditions(*args), 1)} ст. м3/ч\n'

        txtwin.insert(INSERT, result)

    def mass_rate(gas, alpha, f_area, p1):
        k = 1.30 #адиабат. коэффициент для метана
        b3 = 1.59 * ((k / (k + 1)) ** 0.5) * ((2 / (k + 1)) ** (1 / (k - 1))) #расчёт b3

        return 3.16 * b3 * alpha * f_area * ((p1 + 0.1) * gas.actual_density) ** 0.5

    def normal_conditions(gas, alpha, f_area, p1):
        return mass_rate(gas, alpha, f_area, p1) / gas.normal_density

    def standard_conditions(gas, alpha, f_area, p1):
        return normal_conditions(gas, alpha, f_area, p1) * 293.15 / 273.15

    ppk_root = Tk()
    ppk_root.geometry("1920x1080")
    ppk_root.title('Расчёт предохранительных клапанов ГРС ')

    manual_input = ['Давление до клапана, МПа', 'Температура, °С', 'Альфа', 'Расход газа, ст. м3/ч']

    lbl = []
    txt = []

    for i in range(len(manual_input)):
        lbl.append(Label(ppk_root, text=manual_input[i], font=("Arial Bold", 14)))
        lbl[i].grid(column=100, row=700 + 100 * i)

        entry = Entry(ppk_root, width=10, justify='center')
        entry.grid(column=110, row=700 + 100 * i)
        txt.append(entry)

    txtwin = scrolledtext.ScrolledText(ppk_root, width=60, height=10, bg="darkgreen", fg='white')
    txtwin.grid(column=10, row=10)

    combo = Combobox(ppk_root)
    combo['values'] = get_grs_name()
    combo.current(0)
    combo.grid(column=300, row=0)

    btn = Button(ppk_root, text='Выполнить расчёт', font=("Arial Bold", 10), bg="green", fg="black", command=calc)
    btn.grid(column=600, row=100)
    btn1 = Button(ppk_root, text='Расчёт седла', font=("Arial Bold", 10), bg="green", fg="black", command=shift)
    btn1.grid(column=600, row=200)
    btn2 = Button(ppk_root, text='Очистить окно', font=("Arial Bold", 10), bg="green", fg="black", command=clean)
    btn2.grid(column=600, row=300)


    ppk_root.bind('<Return>', calc)

    ppk_root.mainloop()
    return


# ОКНО --- Теплотехнический расчёт, подбор подогревателя газа
def gas_heat_calc():
    heat_root = Tk()
    heat_root.geometry("1920x1080")
    heat_root.title('Теплотехнический расчёт, подбор подогревателя газа ')

    def clean():
        txtwin.delete(1.0, END)  # чистит окно вывода

    def calc():
        """Начнём расчет с нахождения теплоёмкости газа"""
        clean()
        chois = get_composition(combo.get())
        result = (f'\nЗапуск расчёта ГРС {format(combo.get())}\n')
        txtwin.insert(INSERT, result)
        res = [format(data.get()) for data in txt]

        rate = float(res[0].replace(',', '.'))
        inlet_pressure = float(res[1].replace(',', '.'))
        outlet_pressure = float(res[2].replace(',', '.'))
        inlet_temperature = float(res[3].replace(',', '.'))
        grs_pressure_drop = inlet_pressure - outlet_pressure
        outlet_temperature_limit = float(res[4].replace(',', '.'))

        interpolation_coefficient = [-0.06945, -6.3116, 0.0821]  # эти коэффициенты ты получил интерполяцией ВНИИМБ и ТОПКИ и нашел среднее значение перепада температур в заданном диапазоне

        outlet_temperature = inlet_temperature + interpolation_coefficient[0] * grs_pressure_drop**2 + interpolation_coefficient[1] * grs_pressure_drop + interpolation_coefficient[2]

        if outlet_temperature > outlet_temperature_limit:
            outlet_temperature_limit = outlet_temperature


        inlet_gas = Gas(chois, inlet_temperature, inlet_pressure, rate)
        outlet_gas_limit = Gas(chois, outlet_temperature_limit, outlet_pressure, rate)
        outlet_gas = Gas(chois, outlet_temperature, outlet_pressure, rate)

        print(inlet_gas.heat_stream)
        print(outlet_gas.heat_stream)

        heaters = {'ГПМ-ПТПГ-5': 170, 'ГПМ-ПТПГ-10': 300, 'ГПМ-ПТПГ-15М': 500, 'ГПМ-ПТПГ-30М': 1080, 'ГПМ-ПТПГ-100':2700} #это просто каталог доступных подогревателей
        heat_power = (outlet_gas_limit.heat_stream - outlet_gas.heat_stream) / 3600 / 1000

        heater_chois = None

        if outlet_temperature_limit == outlet_temperature:
            heater_chois = "\nПодогреватель газа не требуется!\n"
        else:
            for heater, power in heaters.items():
                print(f'мощность {heater} {power} кВт, а надо {round(heat_power, 2)}')
                if power >= heat_power:
                    heater_chois = f'\nПодобран подогреватель {heater}\n'
                    print(heater_chois)
                    break
            if heater_chois == None:
                heater_chois = ' афигенной мощности'

        """исходные данные для расчёта"""

        txtwin.insert(INSERT,
                      f'\nРасход газа {inlet_gas.rate} ст. м3/ч\n'
                        f'Давление газа на входе ГРС {round((inlet_gas.pressure-101325)/1000000, 2)} МПа\n'
                        f'Температура газа на входе ГРС {round(inlet_gas.temperature-273.15, 2)} °C\n'
                        f'Давление газа на выходе с ГРС {round((outlet_gas.pressure-101325)/1000000, 2)} МПа\n'
                        f'Температура газа на выходе с ГРС без подогрева {round(outlet_gas.temperature-273.15, 2)} °C\n'
                        f'Температура газа на выходе с ГРС не ниже {outlet_temperature_limit} °С\n')

        txtwin.insert(INSERT,
                      f'\nМолярная масса газа - {round(inlet_gas.molecular_mass, 2)} кг/кмоль\n'
                           f'Молярный объём  {round(inlet_gas.molecular_mass / inlet_gas.standard_density, 2)} м3/кмоль\n'
                           f'Массовый расход газа {round(inlet_gas.mass_flow, 1)} кг/ч\n'
                           f'Теплоёмкость газа на входе {round(inlet_gas.specific_heat/1000, 3)} кДж/кг °C\n'
                           f'Теплоёмкость газа на выходе {round(outlet_gas.specific_heat/1000, 3)} кДж/кг °C\n')

        result = (f'Температура газа на выходе ГРС без подогрева газа {round(outlet_temperature, 2)} °С\n'
                  f'требуемая мощность подогревателя  {round(heat_power, 1)} кВт \n'
                  f'{heater_chois}')

        txtwin.insert(INSERT, result)

    manual_input = ['Расход газа, ст. м3/ч',
                    'Давление на входе ГРС, МПа',
                    'Давление на выходе ГРС, МПа',
                    'Температура газа на входе ГРС, °С',
                    'Минимальная тем-ра на выходе, °С']

    lbl = []
    txt = []

    for i in range(len(manual_input)):
        lbl.append(Label(heat_root, text=manual_input[i], font=("Arial Bold", 14)))
        lbl[i].grid(column=100, row=700 + 100 * i)

        entry = Entry(heat_root, width=10, justify='center')
        entry.grid(column=110, row=700 + 100 * i)
        txt.append(entry)

    #entry[4].set(str(-10))

    txtwin = scrolledtext.ScrolledText(heat_root, width=70, height=30, bg="darkgreen", fg='white')
    txtwin.grid(column=10, row=10)

    combo = Combobox(heat_root)
    combo['values'] = get_grs_name()
    combo.current(0)
    combo.grid(column=300, row=0)

    interface_buttons = ['Расчёт', 'Очистить окно']
    calc_list = [calc, clean]

    for i in range(len(interface_buttons)):
        btn = Button(heat_root, text=interface_buttons[i], font=("Arial Bold", 10), bg="green", fg="black",
                     command=calc_list[i])
        btn.grid(column=600, row=100 + 100 * i)

    heat_root.mainloop()
    return


def pipeshell_calc():
    return


def stat_grs():
    grs_stat.make_stat_root()


# ОКНО --- Секции ГРС для расчёта
def insert_tvps_section():
    root = Tk()
    root.geometry('800x600')
    root.title('Секции ГРС для расчёта')

    grs_cmb = Combobox(root, width=30)
    grs_cmb.pack()
    grs_cmb['value'] = get_grs_name()
    grs_cmb.current(0)

    block_cmb = Combobox(root, width=30)
    block_cmb.pack(pady=5)
    block_cmb['values'] = get_column_data(column_name='name_block', table='blocks', name_grs=grs_cmb.get())
    block_cmb.current(0)

    """обработчик событий, вытаскивает все существующие блоки из выбранной на данный момент ГРС"""
    def block_cmb_event(event):
        block_cmb['values'] = get_column_data(column_name='name_block', table='blocks', name_grs=grs_cmb.get())
        block_cmb.current(0)
    grs_cmb.bind("<<ComboboxSelected>>", block_cmb_event)

    data_list = ['Наименование участка', 'Наружний диаметр трубопровода', 'Толщина стенки',
                 'Перепад давления, кПа']
    ent = [Entry(root) for data in data_list]
    lbl = [Label(root, text=data) for data in data_list]
    for i in range(len(ent)):
        lbl[i].pack()
        ent[i].pack(pady=5)
    """функция вставляет секцию для расчёта ТВПС ГРС"""
    def insert():
        connection = sqlite3.connect('grs_database.db')
        cur = connection.cursor()

        section = ent[0].get()
        diameter = int(ent[1].get())
        wall = int(ent[2].get())
        if ent[3].get() == '':
            pressure_drop = 0
        else:
            pressure_drop = float(ent[3].get())

        cur.execute(f"""INSERT INTO tvps(block_id,  section_name, pipe_diameter, wall, internal_diameter, pressure_drop)
                    VALUES(
                    (SELECT block_id FROM blocks WHERE grs_id = (SELECT grs_id FROM grs WHERE name_grs = '{grs_cmb.get()}') 
                    AND name_block = '{block_cmb.get()}'),
                    '{section}',
                    {diameter},
                    {wall},
                    {diameter - 2 * wall}, 
                    {pressure_drop});""")
        connection.commit()
        connection.close()

    insert_btn = Button(root, text='Вставить секцию', bg='green', fg='black', command=insert)
    insert_btn.pack(pady=10)

    root.mainloop()


def get_range(start, stop, step):
    if step == 0:
        print("Нулевой шаг детектед!")
        pass
    else:
        i = start
        while i <= stop:
            yield i
            i += step


"""функция решения СЛАУ методом релаксации, пригодится в дальнейшем, когда появятся СЛАУ"""
def relaxation(coefficients, constants, initial_guess, relaxation_factor, tolerance=1e-6, max_iterations=100):
    n = len(coefficients)
    x = initial_guess #вот этот кусок не нравится, каждый раз вводить изначальный список начальных приближений
    x_new = [0, 0, 0]

    for _ in range(max_iterations):
        for i in range(n):
            sigma = 0
            for j in range(n):
                if j != i:
                    sigma += coefficients[i][j] * x_new[j]
            x_new[i] = (1 - relaxation_factor) * x[i] + (relaxation_factor / coefficients[i][i]) * (constants[i] - sigma)
        if sum([((q - p) ** 2) for q, p in zip(x, x_new)]) ** 0.5 < tolerance:
            print('Нужная точность достигнута')
            return [round(elem, int(- math.log(tolerance, 10))) for elem in x_new]
        x = x_new.copy()
    print('Максимальное количество итераций достигнуто')
    return [round(elem, int(- math.log(tolerance, 10))) for elem in x_new]


# ОКНО --- Расчёт ТВПС
def tvps_calc():
    database_path = 'grs_database.db'

    root = Tk()
    root.geometry('1920x1080')
    root.title('Расчёт ТВПС')

    grs_cmb = Combobox(root, width=30)
    grs_cmb.grid(column=0, row=0)
    grs_cmb['value'] = get_grs_name()
    grs_cmb.current(0)

    connection = sqlite3.connect(database_path)
    cur = connection.cursor()

    section_cmb = Combobox(root, width=40)
    section_cmb.grid(column=0, row=100)
    cur.execute(f"""SELECT section_name FROM tvps WHERE block_id IN (
    SELECT block_id FROM blocks WHERE grs_id = (SELECT grs_id FROM grs WHERE name_grs = '{grs_cmb.get()}'))""")
    sections = [data[0] for data in cur.fetchall()]
    section_cmb['values'] = sections
    if len(section_cmb['values']) != 0:
        section_cmb.current(0)
    else:
        section_cmb.set('СЕКЦИИ НЕ ВСТАВЛЕНЫ!')

    def on_combobox_select(event):
        connection = sqlite3.connect(database_path)
        cur = connection.cursor()
        cur.execute(f"""SELECT section_name FROM tvps WHERE block_id IN (
    SELECT block_id FROM blocks WHERE grs_id = (SELECT grs_id FROM grs WHERE name_grs = '{grs_cmb.get()}'))""")
        sections = [data[0] for data in cur.fetchall()]
        section_cmb['values'] = sections
        if len(section_cmb['values']) != 0:
            section_cmb.current(0)
        else:
            section_cmb.set('СЕКЦИИ НЕ ВСТАВЛЕНЫ!')

    grs_cmb.bind('<<ComboboxSelected>>', on_combobox_select)


    manual_input = ['Минимальное давление', 'Максимальное давление',
                    'Шаг по давлению', 'Минимальная температура',
                    'Максимальная температура', 'Число точек по температуре', 'Масштаб подписей данных']

    lbl = []
    txt = []

    for i in range(len(manual_input)):
        lbl.append(Label(root, text=manual_input[i], font=("Arial Bold", 14)))
        lbl[i].grid(column=500, row=100 * i)

        entry = Entry(root, width=10, justify='center')
        entry.grid(column=600, row=100 * i)
        txt.append(entry)

    def calc():

        connection = sqlite3.connect('grs_database.db')
        cur = connection.cursor()

        res = [float(format(data.get().replace(',', '.'))) for data in txt[:6]]
        pressures = list(get_range(*res[:3]))
        res[5] = (res[4] - res[3]) / res[5]
        temperatures = list(get_range(*res[3:]))

        cur.execute(f"""DELETE FROM tvps_db WHERE grs_name = '{grs_cmb.get()}' AND description = '{section_cmb.get()}'""")
        connection.commit()

        cur.execute(f"""SELECT capacity_after FROM grs WHERE name_grs = '{grs_cmb.get()}';""")
        rate = float(cur.fetchone()[0])

        composition = get_composition(grs_cmb.get())
        cur.execute(f"""SELECT section_name, pipe_diameter, wall, internal_diameter, pressure, temperature,
                capacity, description FROM tvps
                WHERE block_id IN (SELECT block_id FROM blocks WHERE 
                grs_id = (SELECT grs_id FROM grs WHERE name_grs = '{grs_cmb.get()}')) """)
        number = len(pressures) * len(temperatures)
        print(f'Количество точек {number}, расчётное время расчёта {number * 0.3 / 60} минут')
        gas = Gas(composition, temperatures[0], pressures[0], rate)
        cur.execute(f"""SELECT pipe_diameter, wall FROM tvps
                WHERE section_name = '{section_cmb.get()}' AND (block_id IN (
                SELECT block_id FROM blocks WHERE grs_id = (
                SELECT grs_id FROM grs WHERE name_grs = '{grs_cmb.get()}')));""")
        pipe = Pipeline(*list(cur.fetchone()))
        count = 0
        if 'Обводн' in section_cmb.get():
            print('Замечаем, что это обводная линия байпаса')
            speed_limit = 50
        elif 'обводн' in section_cmb.get():
                print('Замечаем, что это обводная линия байпаса')
                speed_limit = 50
        else:
            print('Замечаем, что это основная линия')
            speed_limit = 25
        for pressure in pressures:
            for temperature in temperatures:
                try:
                    actual_rate = gas.get_actual_rate(temperature, pressure)
                    tvps = round(speed_limit * rate / (actual_rate / 3600 / pipe.area))
                except ValueError:
                    print('ПРИБЛИЗИТЕЛЬНЫЙ РАСЧЁТ ПО уравнению регрессии')
                    koeff = 273.15 / (temperature + 273.15) * pipe.area / Pipeline(219, 8).area
                    tvps = round((1095.809670 * pressure ** 2 + 30035.230172 * pressure + 4084.938735) * koeff * speed_limit / 25)

                print(f'Давление {pressure} МПа;\n'
                      f'Температура {temperature}°C\n'
                      f'Скорость {actual_rate / 3600 / pipe.area}\n'
                      f'ТВПС {tvps}')

                cur.execute(f"""INSERT INTO tvps_db(grs_name, description, pressure, 
                    temperature, pipe_diameter, wall, tvps)
                    VALUES('{grs_cmb.get()}', '{section_cmb.get()}', {pressure}, 
                    {temperature}, {pipe.diameter}, {pipe.wall}, {tvps}
                    )""")
                connection.commit()
                count += 1
                progress = f'Прогресс {round(100 * count / number, 2)} %\n'
                print(progress)
        connection.close()


    def draw_plot():
        """считываем базу данных"""
        try:
            min_pressure = float(format(txt[0].get().replace(',', '.')))
            max_pressure = float(format(txt[1].get().replace(',', '.')))
        except ValueError:
            min_pressure = 0
            max_pressure = 1000

        try:
            plot_text_scale = float(format(txt[6].get().replace(',', '.')))
        except ValueError:
            plot_text_scale = 10
        connection = sqlite3.connect(database_path)
        query_max_t = f"""SELECT pipe_diameter AS 'Диаметр', wall AS 'Стенка', 
        pressure AS 'Давление', temperature AS 'Температура', tvps AS 'ТВПС'
        FROM tvps_db WHERE grs_name = '{grs_cmb.get()}' AND description = '{section_cmb.get()}';"""

        main_df = pd.read_sql(query_max_t, connection)



        max_temperature = main_df['Температура'].max()
        min_temperature = main_df['Температура'].min()

        main_df = main_df.loc[main_df['Давление'] > min_pressure, ['Давление', 'Температура', 'ТВПС']]
        main_df = main_df.loc[main_df['Давление'] < max_pressure, ['Давление', 'Температура', 'ТВПС']]

        main_df.sort_values(by=['ТВПС'])

        df = main_df.loc[main_df['Температура'] == max_temperature, ['Давление', 'ТВПС']]

        plt.figure(figsize=(16, 8))

        #plt.plot(df_try['Давление'], df_try['ТВПС'], marker='o', linestyle='-', color='green', markersize=3)

        plt.plot(df['Давление'], df['ТВПС'], marker='o', linestyle='-', color='red', markersize=3,
                 label=f'При температуре газа {round(max_temperature)}°C')
        num_labels = 12
        for i, (x, y) in enumerate(zip(df['Давление'], df['ТВПС'])):
            try:
                if i % (len(df['Давление']) // num_labels) == 0:
                    plt.text(x, y - df['ТВПС'].max() / plot_text_scale, y)
            except ZeroDivisionError:
                num_labels -= 1

        df = main_df.loc[main_df['Температура'] == min_temperature, ['Давление', 'ТВПС']]
        plt.plot(df['Давление'], df['ТВПС'], marker='o', linestyle='-', color='blue', markersize=3,
                 label=f'При температуре газа {round(min_temperature)}°C')

        for i, (x, y) in enumerate(zip(df['Давление'], df['ТВПС'])):
            if i % (len(df['Давление']) // num_labels) == 0:
                plt.text(x, y + df['ТВПС'].max() / plot_text_scale, y)


        plt.legend()
        plt.xlabel('Давление, МПа')
        plt.ylabel('ТВПС, ст м3/ч')
        plt.title(f'ТВПС ГРС {grs_cmb.get()}, {section_cmb.get()}')
        plt.grid(linewidth=1, linestyle='-', color='grey')
        plt.show()


    def view_data():
        connection = sqlite3.connect('grs_database.db')
        cur = connection.cursor()
        view_window = Toplevel(root)
        view_window.title('Таблица ТВПС')

        treeview = ttk.Treeview(view_window)
        treeview.pack()

        treeview["columns"] = ("section_name", "pipe_diameter", "wall",
                               "internal_diameter", "pressure", "temperature",
                               "capacity", "description")
        headings = ("Участок трубопровода, мм", "Диаметр трубопровода, мм", "Толщина стенки, мм",
                               "Внутренний диаметр, мм", "Давление, МПа (изб.)", "Температура, °С",
                               "ТВПС, ст. м3/ч", "Комментарий")
        width_list = [10 * len([letter for letter in word]) for word in headings]
        treeview.column("#0", width=0, stretch=NO)
        for i in range(len(treeview["columns"])):
            treeview.column(treeview["columns"][i], anchor=N, width=width_list[i])
        treeview.heading("#0", text="")
        [treeview.heading(data, text=heading) for data, heading in dict(zip(treeview["columns"], headings)).items()]

        cur.execute(f"""SELECT section_name, pipe_diameter, wall, internal_diameter, pressure, temperature,
        capacity, description FROM tvps
                WHERE block_id IN (SELECT block_id FROM blocks WHERE 
                grs_id = (SELECT grs_id FROM grs WHERE name_grs = '{grs_cmb.get()}')) """)
        rows = cur.fetchall()

        for row in rows:
            treeview.insert("", END, text="", values=row)

        connection.close()

    def show_tvps():

        txtwin.delete(1.0, END)

        connection = sqlite3.connect(database_path)
        cur = connection.cursor()

        inlet_pressures = [1.75, 2.93, 3.00, 3.81, 5.39]
        temperatures = [1, 17]
        for pressure in inlet_pressures:
            """Жесть тут порнография конечно, но я очень спешу, некогда нормально задуматься как делать"""
            if section_cmb.get() == 'От фильтров Ф2.1,2':
                pressure -= 0.01
            elif section_cmb.get() == 'От подогревателя ПГА-200':
                pressure -= 0.11
            elif section_cmb.get() == 'Байпас мимо подогревателя':
                pressure -= 0.11
            elif section_cmb.get() == 'После СГ5.1':
                pressure -= 0.18
            elif section_cmb.get() == 'После СГ5.2':
                pressure -= 0.18
            elif section_cmb.get() == 'Высокая сторона л. редуц. Выход 1,2':
                pressure -= 0.18
            elif section_cmb.get() == 'Обвязка нового ПТПГ-15М':
                pressure -= 0.11
            elif section_cmb.get() == 'Высокая сторона л. ред. после ТП':
                pressure -= 0.18


            for temperature in temperatures:

                print(pressure, temperature)


                cur.execute(f"""SELECT description, pressure as 'Давление, МПа (изб.)', temperature as 'Температура, °С', 
                            MIN(tvps) as 'ТВПС, ст. м3/ч' FROM tvps_db
                            WHERE round(pressure, 2) = {round(pressure, 2)} 
                            AND round(temperature, 2)  = {round(temperature, 2)}
                            AND grs_name = '{grs_cmb.get()}'
                            AND description  = '{section_cmb.get()}'
                            GROUP BY description
                            ORDER BY pressure, tvps;""")
                try:
                    query_result = list(cur.fetchone())
                except TypeError:
                    print(f"""SELECT description, pressure as 'Давление, МПа (изб.)', temperature as 'Температура, °С', 
                            MIN(tvps) as 'ТВПС, ст. м3/ч' FROM tvps_db
                            WHERE round(pressure, 2) = {pressure} 
                            AND round(temperature, 2)  = {temperature}
                            AND grs_name = '{grs_cmb.get()}'
                            AND description  = '{section_cmb.get()}'
                            GROUP BY description
                            ORDER BY pressure, tvps;""")
                text_list = ['Секция', 'Давление, МПа','Температура, °С', 'ТВПС, ст м3/ч']

                result = dict(zip(text_list, query_result))
                #result = [round(data, 2) for data in result[1:]].append(section)
                print(result)

                text = f""" Давление {round(pressure, 2)}, темп. {temperature} ТВПС {result['ТВПС, ст м3/ч']} ст. м3/ч\n"""

                txtwin.insert(INSERT, text)


    treeview_btn = Button(root, text="Таблица", command=view_data, fg='white', bg='black')
    treeview_btn.grid(column=0, row=1000)

    insert_tvps_section_btn = Button(root, text='Вставить секцию ТВПС', command=insert_tvps_section, fg='white', bg='green')
    insert_tvps_section_btn.grid(column=0, row=1100)

    calc_btn = Button(root, text="Расчёт", command=calc, fg='white', bg='black')
    calc_btn.grid(column=300, row=1200)

    plot_btn = Button(root, text='График ТВПС', command=draw_plot, fg='white', bg='blue')
    plot_btn.grid(column=0, row=1300)

    show_tvps_btn = Button(root, text="Показать ТВПС", command=show_tvps, fg='black', bg='pink')
    show_tvps_btn.grid(column=0, row=1400)

    txtwin = scrolledtext.ScrolledText(root, width=140, height=10, bg="darkgreen", fg='white')
    txtwin.grid(column=10, row=10)

    root.mainloop()
    pass


# ОСНОВНОЕ ОКНО
def main_window():
    interface_buttons = ['Скорости в трубопроводах', 'Ёмкость одоранта', 'Пропускная способность клапанов',
                         'Расчёт подогревателя газа', 'Расчёт ППК', 'Толщина стенок трубопроводов', 'Схема',
                         'Статистика', 'База данных ГРС', 'Расчёт ТВПС']
    calc_windows = [veloсity_calc, odorant_calc, valve_calc,
                    gas_heat_calc, ppk_calc, pipe_wall_root,
                    pipe_draw, stat_grs, grs_data, tvps_calc]
    root = Tk()
    root.geometry("450x180")
    root.title('НеВеста-ГРС')

    row_number = int(len(interface_buttons) / 2)
    for i in range(len(interface_buttons)):
        btn = Button(root, text=interface_buttons[i], font=("Arial Bold", 10), bg="green", fg="black",
                     command=calc_windows[i])
        if i < row_number:
            btn.grid(column=100, row=i*100)
        else:
            btn.grid(column=300, row=i*100-100*row_number)
    root.mainloop()


main_window()
#insert_tvps_section()
def pressure_drop_calc():

    # temperature = float(input('Введите температуру, С ').replace(',', '.'))
    # pressure = float(input('Введите давление, МПа (изб) ').replace(',', '.'))
    temperature = 0
    pressure = 0.6
    rate = 76000
    # использовать словарь для определения состава довольно удобно

    gas = Gas(get_composition(random.choice(get_grs_name())), temperature, pressure, rate)
    print(gas.viscosity)
    # тут собирается стринговый аругмент для работы модуля props
    print(
        f'Плотность газа при температуре {round(gas.temperature - 273.15, 2)} С  и давлении {round((gas.pressure - 101325) / 1e6, 2)} МПа, составила {round(gas.actual_density, 4)} кг/м3')
    print(f'Плотность газа при нормальных условиях составила {round(gas.normal_density, 4)} кг/м3')
    print(f'Плотность газа при стандартных условиях составила {round(gas.standard_density, 4)} кг/м3')
    print(f'Теплоёмкость компонентов газа  {gas.component_specific_heat}')
    print(f'Теплоёмкость  газа  {gas.specific_heat}')

    """делай всё это через базы данных SQL пожалста"""
    # pipe_diameter = int(input('Введите внешний диаметр трубопровода, мм '))
    # wall = int(input('Введите толщину стенки трубопровода, мм  '))
    # pipe_diameter = [159, 159, 159, 108, 108, 159, 159, 159, 159, 159, 159, 219, 325, 325, 325, 325, 325]
    # wall = [6, 6, 6, 5, 5, 6, 6, 6, 6, 6, 6, 7, 10, 10, 10, 12, 10]
    # pipe_lenght = [3.5, 3.4, 4.8, 1.7, 2.5, 0.6, 0.6, 1.0, 23.50, 1.3, 3.0, 0.8, 2.3, 30.50, 2.8, 1.2, 9.0]

    pipe_diameter = 325
    wall = 10
    pipe_lenght = 10

    pipe = Pipeline(pipe_diameter, wall, pipe_lenght)

    # rate = float(input('Введите расход газа ст. м3/ч '))

    actual_flow = gas.standard_density / gas.actual_density * rate
    velosity = actual_flow / pipe.area / 3600

    reynolds = velosity * pipe.internal_diameter / 1000 * gas.actual_density / gas.viscosity
    print(f'Число Рейнольдса составило {round(reynolds, 1)}')
    if reynolds > 10000:
        darsi = 0.2579 / (reynolds ** 0.231)
        print(f'коэффицент Дарси при турбулентом движении {round(darsi, 5)}')
    else:
        darsi = 64 / reynolds
        print(f'невероятно, ламинарное течение, что за условия такие {darsi}')
    pressure_drop = darsi * pipe_lenght / pipe_diameter * 1000 * velosity ** 2 * gas.actual_density / 2

    # final_pressure = 1e6 * float(input('Введите конечное давление, МПа ').replace(',', '.')) + 101325
    # gas.expansion(final_pressure)

    # total_specific_heat = sum({key: (gas.specific_heat[key]*gas.mass_fraction[key]) for key in gas.composition}.values())
    # print(total_specific_heat) #типа полная теплоёмкость газа
    # print(gas.entalpy)
